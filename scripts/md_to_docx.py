import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Skip mermaid blocks and insert actual image
        if line.strip().startswith('```mermaid'):
            while i < len(lines) and not lines[i].strip() == '```':
                i += 1
            i += 1
            # Insert the generated flowchart image
            img_path = r"C:\Users\asr26\.gemini\antigravity\brain\c8a83804-e778-45e8-925b-33afa097b741\pipeline_flowchart_1768903229692.png"
            try:
                doc.add_picture(img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f'[Diagram could not be loaded: {e}]')
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            i += 1
            continue
        
        # Tables
        if line.strip().startswith('|') and '|' in line:
            table_rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            i += 1
            continue
        elif table_rows:
            # End of table, render it
            if len(table_rows) > 1:
                # Skip separator row
                header = table_rows[0]
                data = [row for row in table_rows[1:] if not all(c.startswith('-') for c in row)]
                
                if data:
                    table = doc.add_table(rows=1 + len(data), cols=len(header))
                    table.style = 'Table Grid'
                    
                    # Header
                    hdr_cells = table.rows[0].cells
                    for j, h in enumerate(header):
                        hdr_cells[j].text = h
                        for run in hdr_cells[j].paragraphs[0].runs:
                            run.bold = True
                    
                    # Data rows
                    for row_idx, row_data in enumerate(data):
                        row_cells = table.rows[row_idx + 1].cells
                        for j, cell_text in enumerate(row_data):
                            if j < len(row_cells):
                                row_cells[j].text = cell_text
                    
                    doc.add_paragraph()
            table_rows = []
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        elif line.strip().startswith('- '):
            # Bullet
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif line.strip().startswith('1. '):
            p = doc.add_paragraph(line.strip()[3:], style='List Number')
        elif line.strip():
            # Regular paragraph - handle bold/links
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold markers (will add back)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links, keep text
            doc.add_paragraph(text)
        
        i += 1
    
    # Handle any remaining table
    if table_rows and len(table_rows) > 1:
        header = table_rows[0]
        data = [row for row in table_rows[1:] if not all(c.startswith('-') for c in row)]
        if data:
            table = doc.add_table(rows=1 + len(data), cols=len(header))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for j, h in enumerate(header):
                hdr_cells[j].text = h
            for row_idx, row_data in enumerate(data):
                row_cells = table.rows[row_idx + 1].cells
                for j, cell_text in enumerate(row_data):
                    if j < len(row_cells):
                        row_cells[j].text = cell_text

    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

if __name__ == "__main__":
    md_to_docx(
        r"C:\Users\asr26\.gemini\antigravity\brain\c8a83804-e778-45e8-925b-33afa097b741\pipeline_walkthrough.md",
        r"C:\Users\asr26\.gemini\antigravity\brain\c8a83804-e778-45e8-925b-33afa097b741\pipeline_walkthrough.docx"
    )
